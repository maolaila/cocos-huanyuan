from __future__ import annotations

import json
import math
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = PROJECT_ROOT / "docs"
OUTPUT_ROOT = PROJECT_ROOT / "output"
PLAYWRIGHT_ROOT = OUTPUT_ROOT / "playwright"
DOCUMENT_ASSET_ROOT = OUTPUT_ROOT / "document-assets"
DOCUMENT_ASSET_ROOT.mkdir(parents=True, exist_ok=True)

OUTPUT_DOCX = DOCS_ROOT / "dzpk-phase-a-original-cocos-study-guide.docx"
ACCEPTANCE_JSON = DOCS_ROOT / "evidence" / "phase-a-acceptance-final.json"

FONT_CJK = "Microsoft YaHei"
FONT_CODE = "Consolas"
FONT_FILE = Path("C:/Windows/Fonts/msyh.ttc")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 39, 58)
MUTED = RGBColor(90, 100, 112)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RISK_FILL = "FDECEC"
POSITIVE_FILL = "EAF5EA"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name=FONT_CJK, size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must total {TABLE_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_pr.append(tbl_ind)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    margins = OxmlElement("w:tblCellMar")
    for side, width in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(width))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tbl_pr.append(margins)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[column_index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[column_index]))


def paragraph_border_bottom(paragraph, color="D8DEE8", size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = document.styles["Normal"]
    normal.font.name = FONT_CJK
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_CJK)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CJK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (11.5, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = FONT_CJK
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_CJK)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CJK)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = document.styles[list_style_name]
        style.font.name = FONT_CJK
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.22

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_paragraph.add_run("KG DZPK | Creator 2.4.7 原版还原学习手册")
    set_run_font(header_run, size=8.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_paragraph.add_run("Phase A Creator 2.4.7 Parity Verified  |  ")
    set_run_font(footer_run, size=8.5, color=MUTED)
    add_page_field(footer_paragraph, "PAGE")
    separator = footer_paragraph.add_run(" / ")
    set_run_font(separator, size=8.5, color=MUTED)
    add_page_field(footer_paragraph, "NUMPAGES")

    props = document.core_properties
    props.title = "KG 德州扑克 Phase A 原版 Cocos 学习手册"
    props.subject = "Creator 2.4.7 原 Prefab、语义化源码与 GameHub TRIAL Authority"
    props.author = ""
    props.keywords = "Cocos Creator, DZPK, Texas Holdem, GameHub, Prefab, WebSocket"


def add_paragraph(document, text="", bold_prefix=None, italic=False, color=None, align=None, after=6):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True, color=color or INK)
        remainder = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(remainder, italic=italic, color=color or INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic, color=color or INK)
    return paragraph


def add_heading(document, text, level=1):
    return document.add_heading(text, level=level)


def add_bullets(document, entries):
    for entry in entries:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(entry)
        set_run_font(run, size=10.5)


def create_restarted_list_number(document):
    """Reuse Word's built-in List Number definition with a fresh start at 1."""
    list_style = document.styles["List Number"]._element
    style_p_pr = list_style.find(qn("w:pPr"))
    style_num_pr = style_p_pr.find(qn("w:numPr")) if style_p_pr is not None else None
    style_num_id = style_num_pr.find(qn("w:numId")) if style_num_pr is not None else None
    if style_num_id is None:
        raise RuntimeError("List Number style has no numbering definition")

    numbering = document.part.numbering_part.element
    base_num_id = style_num_id.get(qn("w:val"))
    base_num = next(
        (
            element
            for element in numbering.findall(qn("w:num"))
            if element.get(qn("w:numId")) == base_num_id
        ),
        None,
    )
    if base_num is None:
        raise RuntimeError(f"List Number numId {base_num_id} is missing")
    abstract_reference = base_num.find(qn("w:abstractNumId"))
    if abstract_reference is None:
        raise RuntimeError("List Number abstract definition is missing")

    existing_num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    restarted_num_id = max(existing_num_ids, default=0) + 1
    number_instance = OxmlElement("w:num")
    number_instance.set(qn("w:numId"), str(restarted_num_id))
    new_abstract_reference = OxmlElement("w:abstractNumId")
    new_abstract_reference.set(qn("w:val"), abstract_reference.get(qn("w:val")))
    number_instance.append(new_abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    number_instance.append(level_override)
    numbering.append(number_instance)
    return restarted_num_id


def add_numbered(document, entries):
    restarted_num_id = create_restarted_list_number(document)
    for entry in entries:
        paragraph = document.add_paragraph(style="List Number")
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        level_reference = OxmlElement("w:ilvl")
        level_reference.set(qn("w:val"), "0")
        number_reference = OxmlElement("w:numId")
        number_reference.set(qn("w:val"), str(restarted_num_id))
        num_pr.extend([level_reference, number_reference])
        run = paragraph.add_run(entry)
        set_run_font(run, size=10.5)


def add_code_block(document, code_text):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    for line_index, line in enumerate(code_text.strip("\n").splitlines()):
        if line_index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name=FONT_CODE, size=8.4, color=RGBColor(36, 41, 47))
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(document, label, text, fill=CALLOUT, accent=BLUE):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(label + "  ")
    set_run_font(label_run, bold=True, color=accent, size=10.5)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=INK, size=10.5)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(document, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=9.0):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for column_index, header in enumerate(headers):
        cell = header_row.cells[column_index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=DARK_BLUE)
    for row_values in rows:
        row = table.add_row()
        for column_index, value in enumerate(row_values):
            paragraph = row.cells[column_index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if column_index == 0 and len(headers) > 2
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    # Rows are appended after the first geometry pass, so apply it once more to
    # give every body row fixed widths and w:cantSplit page-break protection.
    set_table_geometry(table, widths_dxa)
    document.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_figure(document, image_path, caption, width_inches=6.35):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(2)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    run = caption_paragraph.add_run(caption)
    set_run_font(run, size=8.5, italic=True, color=MUTED)


def load_font(size):
    if FONT_FILE.exists():
        return ImageFont.truetype(str(FONT_FILE), size=size)
    return ImageFont.load_default()


def wrap_draw_text(draw, box, text, font, fill, max_chars=20, align="center"):
    lines = []
    for source_line in text.split("\n"):
        lines.extend(textwrap.wrap(source_line, width=max_chars) or [""])
    line_height = font.size + 10 if hasattr(font, "size") else 28
    total_height = line_height * len(lines)
    y = box[1] + (box[3] - box[1] - total_height) / 2
    for line in lines:
        bounds = draw.textbbox((0, 0), line, font=font)
        width = bounds[2] - bounds[0]
        x = box[0] + 18 if align == "left" else box[0] + (box[2] - box[0] - width) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def draw_box(draw, box, title, subtitle, fill, border, title_color=(255, 255, 255)):
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=border, width=4)
    title_font = load_font(38)
    subtitle_font = load_font(31)
    title_box = (box[0], box[1] + 6, box[2], box[1] + 64)
    wrap_draw_text(draw, title_box, title, title_font, title_color, max_chars=30)
    subtitle_box = (box[0] + 18, box[1] + 62, box[2] - 18, box[3] - 8)
    wrap_draw_text(draw, subtitle_box, subtitle, subtitle_font, (33, 47, 61), max_chars=52)


def draw_arrow(draw, start, end, color=(64, 91, 123), width=6):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_length = 20
    points = [
        end,
        (
            end[0] - head_length * math.cos(angle - math.pi / 6),
            end[1] - head_length * math.sin(angle - math.pi / 6),
        ),
        (
            end[0] - head_length * math.cos(angle + math.pi / 6),
            end[1] - head_length * math.sin(angle + math.pi / 6),
        ),
    ]
    draw.polygon(points, fill=color)


def create_architecture_diagram(output_path):
    canvas = Image.new("RGB", (1900, 1120), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((65, 35), "DZPK Phase A - 原版 Cocos 与 GameHub Authority 架构", font=load_font(48), fill=(31, 77, 120))
    layers = [
        ((80, 135, 1820, 315), "原版可见层", "Load / Room / DZPKMain / Rule / Set / Night\n原节点、Atlas、Spine、Particle、Animation", (221, 235, 247)),
        ((80, 365, 1820, 545), "语义化 Cocos 层", "5 个语义核心组件 + Standalone 服务\n直接挂载原 Prefab，不重画牌桌", (226, 242, 232)),
        ((80, 595, 1820, 775), "协议与会话层", "context/init + KG Base64 WS + viewer projection\nMsg_Hall_* / Msg_DZPK_*", (255, 244, 214)),
        ((80, 825, 1820, 1035), "GameHub 服务端权威层", "Redis CAS/AOF + engine + bot + control\n完整 deal、私牌、赢家和余额只在服务端", (247, 226, 226)),
    ]
    for box, title, subtitle, fill in layers:
        draw_box(draw, box, title, subtitle, fill, (46, 116, 181), title_color=(31, 77, 120))
    for first, second in zip(layers, layers[1:]):
        start = ((first[0][0] + first[0][2]) / 2, first[0][3] + 8)
        end = ((second[0][0] + second[0][2]) / 2, second[0][1] - 8)
        draw_arrow(draw, start, end)
    canvas.save(output_path)


def create_sequence_diagram(output_path):
    canvas = Image.new("RGB", (1900, 1080), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((60, 30), "原版运行时事件时序", font=load_font(48), fill=(31, 77, 120))
    participants = ["Cocos Boot", "原 Prefab", "GameHub API", "KG WS Adapter", "Redis/Engine"]
    x_positions = [170, 520, 900, 1280, 1650]
    header_font = load_font(27)
    body_font = load_font(24)
    for name, x in zip(participants, x_positions):
        bounds = draw.textbbox((0, 0), name, font=header_font)
        draw.rounded_rectangle((x - 125, 95, x + 125, 155), radius=12, fill=(232, 238, 245), outline=(46, 116, 181), width=3)
        draw.text((x - (bounds[2] - bounds[0]) / 2, 108), name, font=header_font, fill=(31, 77, 120))
        draw.line((x, 155, x, 1015), fill=(180, 190, 202), width=3)
    steps = [
        (170, 900, 210, "POST context/init\nlaunchCode -> sessionToken"),
        (170, 1280, 315, "WS Msg_Hall_Connect\n绑定 source uid"),
        (170, 520, 420, "instantiate Load\nGameSessions -> Room"),
        (520, 1280, 525, "EnterRoom / FinishLoad"),
        (1280, 1650, 630, "CAS restore/create\ncommit complete deal"),
        (1280, 520, 735, "RoomInfo -> FaCards -> StageBet\nCallUserAct / ActBet / PublicCards"),
        (1280, 520, 840, "Result (15s)\nwinners / pots / returns"),
        (1650, 1280, 945, "SETTLED -> WAITING -> PREPARING\n服务端自动下一手"),
    ]
    for start_x, end_x, y, label in steps:
        draw_arrow(draw, (start_x, y), (end_x, y), color=(64, 91, 123), width=5)
        mid_x = (start_x + end_x) / 2
        lines = label.split("\n")
        text_y = y - 58
        for line in lines:
            bounds = draw.textbbox((0, 0), line, font=body_font)
            draw.rectangle((mid_x - (bounds[2] - bounds[0]) / 2 - 8, text_y - 2, mid_x + (bounds[2] - bounds[0]) / 2 + 8, text_y + 30), fill="white")
            draw.text((mid_x - (bounds[2] - bounds[0]) / 2, text_y), line, font=body_font, fill=(33, 47, 61))
            text_y += 30
    canvas.save(output_path)


def create_state_diagram(output_path):
    canvas = Image.new("RGB", (1900, 820), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((60, 30), "服务端 13 状态与恢复边界", font=load_font(48), fill=(31, 77, 120))
    states = [
        "WAITING", "PREPARING", "HOLD_READY", "DEAL_COMMITTED", "DEALING",
        "PREFLOP", "FLOP", "TURN", "RIVER", "SHOWDOWN", "TERMINAL", "SETTLING", "SETTLED",
    ]
    columns = 5
    box_width = 300
    box_height = 95
    x_gap = 55
    y_gap = 105
    start_x = 90
    start_y = 145
    centers = []
    for index, state in enumerate(states):
        row = index // columns
        column = index % columns
        x = start_x + column * (box_width + x_gap)
        y = start_y + row * (box_height + y_gap)
        fill = (226, 242, 232) if state in {"WAITING", "SETTLED"} else (232, 238, 245)
        draw.rounded_rectangle((x, y, x + box_width, y + box_height), radius=18, fill=fill, outline=(46, 116, 181), width=3)
        bounds = draw.textbbox((0, 0), state, font=load_font(25))
        draw.text((x + (box_width - (bounds[2] - bounds[0])) / 2, y + 31), state, font=load_font(25), fill=(31, 77, 120))
        centers.append((x + box_width / 2, y + box_height / 2))
    for index in range(len(states) - 1):
        current = centers[index]
        following = centers[index + 1]
        if index % columns == columns - 1:
            draw_arrow(draw, (current[0], current[1] + box_height / 2), (following[0], following[1] - box_height / 2))
        else:
            draw_arrow(draw, (current[0] + box_width / 2, current[1]), (following[0] - box_width / 2, following[1]))
    draw_arrow(draw, (centers[-1][0], centers[-1][1] + 55), (centers[0][0], centers[0][1] + 55), color=(155, 93, 27), width=5)
    draw.text((665, 720), "每个状态均可版本化快照恢复；SETTLED -> WAITING 已回测", font=load_font(27), fill=(122, 90, 0))
    canvas.save(output_path)


def create_collage(output_path, entries, title, columns=2, cell_width=850, cell_height=520):
    rows = math.ceil(len(entries) / columns)
    canvas = Image.new("RGB", (columns * cell_width + 80, rows * (cell_height + 55) + 120), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((40, 25), title, font=load_font(44), fill=(31, 77, 120))
    label_font = load_font(27)
    for index, (label, image_path) in enumerate(entries):
        row = index // columns
        column = index % columns
        x = 40 + column * cell_width
        y = 95 + row * (cell_height + 55)
        image = Image.open(image_path).convert("RGB")
        image.thumbnail((cell_width - 35, cell_height - 45), Image.Resampling.LANCZOS)
        paste_x = x + (cell_width - image.width) // 2
        paste_y = y + 42 + (cell_height - 45 - image.height) // 2
        canvas.paste(image, (paste_x, paste_y))
        draw.text((x + 10, y), label, font=label_font, fill=(33, 47, 61))
        draw.rectangle((x, y + 35, x + cell_width - 20, y + cell_height), outline=(210, 218, 228), width=2)
    canvas.save(output_path, quality=92)


def build_visual_assets():
    architecture = DOCUMENT_ASSET_ROOT / "architecture.png"
    sequence = DOCUMENT_ASSET_ROOT / "runtime-sequence.png"
    state = DOCUMENT_ASSET_ROOT / "state-machine.png"
    create_architecture_diagram(architecture)
    create_sequence_diagram(sequence)
    create_state_diagram(state)

    ui_flow = DOCUMENT_ASSET_ROOT / "ui-flow-collage.jpg"
    create_collage(ui_flow, [
        ("原 Room.prefab", PLAYWRIGHT_ROOT / "final-room.png"),
        ("原 DZPKMain.prefab", PLAYWRIGHT_ROOT / "final-table.png"),
        ("River / 原牌面高亮", PLAYWRIGHT_ROOT / "board-5-cards.png"),
        ("自然 Result / pot rows", PLAYWRIGHT_ROOT / "result-hand-1.png"),
    ], "原版可见运行链", columns=2)

    popup_flow = DOCUMENT_ASSET_ROOT / "popup-flow-collage.jpg"
    create_collage(popup_flow, [
        ("Rule", PLAYWRIGHT_ROOT / "final-rule.png"),
        ("Set", PLAYWRIGHT_ROOT / "final-set.png"),
        ("Set - Night", PLAYWRIGHT_ROOT / "final-set-night.png"),
        ("844 x 390", PLAYWRIGHT_ROOT / "mobile-landscape-844x390.png"),
    ], "弹层、昼夜与移动横屏", columns=2)

    settlement_flow = DOCUMENT_ASSET_ROOT / "settlement-collage.jpg"
    create_collage(settlement_flow, [
        ("自然 Result", PLAYWRIGHT_ROOT / "result-hand-2.png"),
    ], "原版自然结算", columns=1, cell_width=850, cell_height=560)
    return architecture, sequence, state, ui_flow, popup_flow, settlement_flow


def add_cover(document, hero_image):
    for _ in range(3):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("COCOS CREATOR 2.4.7  ·  ORIGINAL CLIENT RESTORATION")
    set_run_font(run, size=10.5, bold=True, color=BLUE)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("KG 德州扑克（DZPK）")
    set_run_font(run, size=29, bold=True, color=DARK_BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    run = subtitle.add_run("原 Prefab · 语义化源码 · GameHub TRIAL Authority")
    set_run_font(run, size=15, color=RGBColor(43, 81, 99))

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(22)
    run = metadata.add_run("Phase A 学习手册  |  Source commit b5694d57  |  2026-08-25")
    set_run_font(run, size=9.5, italic=True, color=MUTED)
    add_figure(document, hero_image, "最终 exact build：原 DZPKMain、六席、原按钮与 viewer-safe 私牌", width_inches=6.2)
    document.add_page_break()


def build_document():
    acceptance = json.loads(ACCEPTANCE_JSON.read_text(encoding="utf-8"))
    architecture, sequence, state_diagram, ui_flow, popup_flow, settlement_flow = build_visual_assets()
    document = Document()
    configure_document(document)
    add_cover(document, PLAYWRIGHT_ROOT / "final-table.png")

    add_heading(document, "目录与使用方法", 1)
    add_bullets(document, [
        "第 1-4 章：先建立工程事实、架构和目录认知。",
        "第 5-8 章：沿原 Prefab、语义脚本、Cocos API 和协议时序读代码。",
        "第 9-12 章：理解规则、机器人、控牌、视觉、重连和安全边界。",
        "第 13-16 章：复盘问题决策、构建测试、面试问答和后续难度。",
        "附录：Class ID、事件、命令、测试数字和简历表述。",
    ])
    add_callout(document, "建议学习顺序", "先在 Creator 打开原 Prefab 看节点和属性，再按事件从 Controller 跟到 Presentation，最后对照 GameHub projection 与 engine。不要先背 API。")

    add_heading(document, "1. 结论、范围与事实边界", 1)
    add_paragraph(document, "本阶段完成的是独立 Creator 2.4.7 原版客户端学习工程：一个真人、五个服务端机器人、完整 TRIAL 牌局、viewer-safe 协议、Redis 快照、刷新/进程重启恢复。GameHub 前端没有修改。")
    add_table(document, ["事实", "当前结论", "边界"], [
        ("原版画面", "Load / Room / DZPKMain / Rule / Set 已实际运行", "原 Prefab 节点与资源"),
        ("核心源码", "5 个语义文件实际挂载/导入", "原 5 脚本保留作字节证据"),
        ("牌局 authority", "GameHub backend + Redis + engine", "客户端只提交动作意图"),
        ("金额模式", "TRIAL only", "REAL 在 context 和 WS 前置拒绝"),
        ("真人数量", "1 真人 + 5 bot", "2-6 真人接口预留，未启用"),
        ("Creator 3.8", "未开始", "属于下一阶段，不冒充完成"),
        ("Bank", "源码保留、运行时 N/A", "不接独立 DZPK Boot"),
    ], [2200, 3980, 3180])
    add_callout(document, "准确表述", "fresh independent review 已给出 Creator247OriginalClientParityVerified。可以说自己完成了原版还原、语义化重构、服务端机器人/快照/控牌和 GameHub TRIAL 接入；不要说已完成真钱、2-6 真人、3.8 或生产上线。", fill="FFF8E8", accent=RGBColor(122, 90, 0))

    add_heading(document, "2. 权威源码与差异审计", 1)
    add_table(document, ["对象", "权威身份", "还原策略"], [
        ("Cocos", r"C:\total\kg-cocos-client\728_mobile_restore @ b5694d57", "只读；完整 DZPK + meta"),
        ("PHP", r"C:\total\kg-php", "只读规则/协议证据，不作为 runtime"),
        ("目标", str(PROJECT_ROOT), "独立 Creator 2.4.x 工程"),
        ("GameHub", r"C:\total\game-hub develop", "只改 backend/shared/engine"),
    ], [1700, 4400, 3260])
    add_paragraph(document, "完整 `assets/DZPK` 为 695 个文件、10,781,546 bytes。六个 Prefab 根尺寸均为 1334×750；DZPKMain 有 321 节点。原主桌含 25 个 Spine、6 个 ParticleSystem、7 个 Animation。")
    add_bullets(document, [
        "SourceDefectRepair：删除两条 target 为空的 ERNN/BJL 跨游戏 ClickEvent 引用。",
        "IntentionalDifference：不恢复 PHP 单赢家、花色破平、伪边池、固定最小加注等规则缺陷。",
        "PlannedLaterPhase：REAL money、正式 UID policy、2-6 真人、Creator 3.8、线上 package。",
    ])

    add_heading(document, "3. 总体架构", 1)
    add_figure(document, architecture, "图 1  原可见层、语义 Cocos、协议会话与服务端 authority 分层")
    add_paragraph(document, "关键设计不是“原脚本一字不改”或“重新画一套”，而是保留原 Prefab/节点/资源，把含糊脚本重写成语义组件，再用原 Class ID 映射位置实际挂载。")
    add_code_block(document, """
Original visible source
  Load.prefab -> Room.prefab -> DZPKMain.prefab
                          +-> Rule.prefab
                          +-> Set.prefab -> Night

Semantic Cocos code
  DzpkLoadingScreenController
  DzpkRoomSelectionController
  DzpkTableStateModel
  DzpkTableGameController
  DzpkTablePresentation

Server authority
  GameHub context/init -> authenticated KG WS
  -> DzpkStudyAuthority -> Redis CAS/AOF -> kg-dzpk-engine
""")

    add_heading(document, "4. 工程目录与启动入口", 1)
    add_code_block(document, """
kg-dzpk-2.4.7-original/
  assets/
    Scene/DzpkStandaloneBoot.fire
    Standalone/                 # Cocos 内置会话/资源/音频/导航能力
    DZPK/
      prefab/                   # 原 6 Prefab
      _script/                  # 原 5 核心脚本 + 证据
      _semantic/                # 当前可维护实现
      AA, _res, sound, ...      # 原资源与动画
    _script/                    # 10 shared + 7 provider
    BJL/.../plist_puke          # 原扑克牌 Atlas
    resources/Hall, sound       # 头像/公共音效
  scripts/
    materialize-original-source.ps1
    apply-semantic-prefab-mapping.ps1
    build-creator247.ps1
    verify-original-client-boundary.ps1
    playwright-*.js
  docs/
""")
    add_numbered(document, [
        "Boot 交换 launchCode，取得 opaque sessionToken；从 URL 清除一次性凭证。",
        "连接 authenticated KG WS，Msg_Hall_Connect 回写服务端 source uid。",
        "加载 DZPK Asset Bundle，再实例化原 Load。",
        "Load 请求 GameSessions，实例化原 Room；玩家点击房间。",
        "Room 请求 EnterRoom，实例化原 DZPKMain；PokerBase 发送 FinishLoad。",
        "RoomInfo 建立 viewer snapshot，后续事件驱动原节点和动画。",
    ])

    add_heading(document, "5. 原 Prefab 与组件映射", 1)
    add_table(document, ["Prefab", "原 UUID / 状态", "当前活跃组件", "作用"], [
        ("Load", "07af51a2...", "41f21+8TLlAnLWc/P9hpsrY", "BGM、Spine、配置、Room"),
        ("Room", "efa3b3c7...", "84ec95q06xDTpmQT8hjRm0+", "三房、头像、进入、Rule"),
        ("DZPKMain Controller", "faea1885...", "6b8f4AeDNVOdaQkSDFsDwLn", "事件、状态、动作编排"),
        ("DZPKMain Presentation", "同一 Prefab", "2de8849XWxHsblCkTlm5qw/", "原节点/牌/筹码/动画"),
        ("Rule", "1eb86003...", "原 Rule 组件", "规则弹层"),
        ("Set", "31f79f27...", "原 Set 组件", "音量、昼夜"),
        ("Bank", "8c9ac144...", "不活跃", "SourcePresentButGameRuntimeNotApplicable"),
    ], [1750, 2300, 2800, 2510], font_size=8.5)
    add_figure(document, ui_flow, "图 2  原 Room、主桌、River 与自然 Result")
    add_paragraph(document, "Prefab 仍是布局与资源事实源；`.meta` UUID 保证序列化资源引用稳定。`apply-semantic-prefab-mapping.ps1` 只替换组件 Class ID、语义属性名和事件 handler，不重排节点。")

    document.add_page_break()
    add_heading(document, "6. 五个语义核心脚本", 1)
    add_table(document, ["脚本", "职责", "关键方法"], [
        ("DzpkLoadingScreenController", "原 Load 生命周期", "initializeOriginalLoadingFlow / requestOriginalRoomConfigurations"),
        ("DzpkRoomSelectionController", "原 Room 选房", "requestOriginalRoomEntry / instantiateOriginalMainTable"),
        ("DzpkTableStateModel", "viewer-safe 表状态", "initializeFromRoomSnapshot / calculate*Presets"),
        ("DzpkTableGameController", "事件订阅与编排", "handleRoomSnapshotReceived / handleHandSettled"),
        ("DzpkTablePresentation", "原节点与动画", "animateHoleCardDeal / animateStandardPotDistribution"),
    ], [3000, 2100, 4260], font_size=8.3)
    add_callout(document, "兼容代理原则", "m_roomInfo、Msg_DZPK_*、onClick 等旧名只保留一行代理；新业务逻辑必须进入语义方法。序列化按钮已改绑语义 handler。")
    add_table(document, ["原方法", "语义方法", "学习重点"], [
        ("initShow", "initializeTablePresentation", "节点缓存、NodePool、初始显隐"),
        ("add_Delete_Player", "setParticipantSeatPresence", "六席本地座位映射"),
        ("setPlayerBet", "animateParticipantWager", "NodePool + cc.Action"),
        ("showFaPai", "animateHoleCardDeal", "两轮发牌、本人翻面"),
        ("showPubliccards", "animateCommunityCardReveal", "flop/turn/river"),
        ("playerBrightCard", "renderParticipantShowdown", "牌型、灰度、winner"),
        ("showGetPlayerBet", "animateStandardPotDistribution", "主池/边池/退款"),
        ("sliderEvevt", "handleRaiseSliderChanged", "Slider/ProgressBar/Spine"),
    ], [2100, 3600, 3660], font_size=8.5)

    add_heading(document, "7. Cocos Creator 技术与 API", 1)
    cocos_rows = [
        ("cc.Component / 生命周期", "onLoad、onEnable、start、onDestroy", "Boot、Load、Room、Table、Set"),
        ("properties", "Prefab 序列化引用", "Seat root、Label、SpriteAtlas、pot node"),
        ("Node / cc.find", "层级查询、active、opacity、color、zIndex", "不重画节点"),
        ("cc.Prefab / instantiate", "Load -> Room -> Main；Rule/Set popup", "Asset Bundle 内实例化"),
        ("Sprite / SpriteFrame", "牌背、牌面、按钮状态", "Atlas frame name 映射"),
        ("SpriteAtlas", "BJL plist_puke + DZPK type atlas", "rank/suit/hand type"),
        ("Button / ClickEvent", "Prefab handler + customEventData", "fold/call/check/raise/menu"),
        ("Toggle", "自动动作、声音、音乐、昼夜", "Set 与 auto actions"),
        ("Slider / ProgressBar", "加注金额与全下", "progress clamp + 31 segments"),
        ("Label / BitmapFont", "余额、底池、派奖", "退款语义放 pot row 避免缺字"),
        ("Animation / Clip", "倒计时、Slider spark", "speed = 1 / seconds"),
        ("Spine", "Load、环境、all-in、winner、Big Win", "保留 animation/start/idle"),
        ("ParticleSystem", "原主桌粒子效果", "Prefab 序列化资源"),
        ("Widget", "1334×750 和 844×390 适配", "SHOW_ALL / 原对齐"),
        ("EventTarget / EventBus", "wGEvent subscribe/publish/off", "目标有效性与事件名不改"),
        ("schedule / cc.Action", "动画等待、移动、淡入淡出", "保留原时序，抑制已知 deprecation"),
        ("Asset Bundle", "cc.assetManager.loadBundle('DZPK')", "脚本和资源按 bundle 加载"),
        ("audioEngine", "原 AudioManager + browser gesture gate", "BGM sound/back"),
        ("Scene / Camera / Canvas", "无可见大厅的独立入口", "容器不是外部网页壳"),
        (".meta / UUID", "资源和组件序列化身份", "还原工程最重要的稳定边界"),
    ]
    add_table(document, ["API/技术", "如何使用", "本工程位置"], cocos_rows, [1900, 3500, 3960], font_size=8.2)

    add_heading(document, "8. 原事件协议与后端交互", 1)
    add_figure(document, sequence, "图 3  从启动凭证到自动下一手的原事件时序")
    add_table(document, ["方向", "事件", "核心字段 / 行为"], [
        ("Cocos -> Server", "Msg_Hall_GameSessions", "gtype=19"),
        ("Cocos -> Server", "Msg_Hall_EnterRoom", "tableid=0, gtype=19, level"),
        ("Cocos -> Server", "Msg_Hall_FinishLoad", "rid"),
        ("Cocos -> Server", "Msg_DZPK_ActBet", "gold；request_id 由 transport 加"),
        ("Cocos -> Server", "Msg_DZPK_Out", "退出意图"),
        ("Server -> Cocos", "Msg_DZPK_RoomInfo", "stage/players/notice/publiccards/pot/px"),
        ("Server -> Cocos", "FaCards / StageBet", "viewer 私牌 / 盲注前注"),
        ("Server -> Cocos", "CallUserAct / ActBet", "uid/minbet/time / actor uid/actionSeq"),
        ("Server -> Cocos", "PublicCards", "cards + viewer-only px"),
        ("Server -> Cocos", "Result", "handId/revision/actionSeq/winners/pots/returns/cards"),
    ], [1600, 3000, 4760], font_size=8.3)
    add_paragraph(document, "动作 gold 映射：负数为 FOLD；0 且 toCall=0 为 CHECK；等于 toCall 为 CALL；等于剩余 stack 为 ALL_IN；其它合法值映射为 raiseTo。身份来自认证 connection，客户端 uid 只做一致性。")
    add_callout(document, "REAL 边界", "Phase A 在 context/init 与 KG WS 两处拒绝 DZPK REAL，早于钱包读取。真人真钱 hold/capture/order/ledger/audit 不在本阶段。", fill=RISK_FILL, accent=RGBColor(155, 28, 28))

    add_heading(document, "9. 标准德州规则与 13 状态", 1)
    add_figure(document, state_diagram, "图 4  13 状态、版本化快照与下一手边界")
    add_bullets(document, [
        "52 张唯一牌；rank*100+suit 编码；5-7 张 best five；A2345 wheel。",
        "花色不破平；完整 kicker chain；同牌型精确比较。",
        "按贡献层构造主池和多个边池；仅 eligible participant 争夺该层。",
        "平分与 odd chip 按庄后顺序确定性分配；未跟注 overbet 原样退回且不抽 rake。",
        "full raise 决定最小再加注；short all-in 不重新开放已行动玩家加注权。",
        "clientActionId/request_id 幂等；expected revision 和 seat/hand 绑定。",
        "每次动作、街道和 13 状态都可恢复；失败写入不应推进 authority。",
    ])
    add_table(document, ["测试", "最终数字", "结论"], [
        ("5-card exhaustive", "2,598,960", "分类计数完整"),
        ("7-card independent oracle", "1,000,000 / 0 mismatch", "best-five 一致"),
        ("Pot/rake conservation", "contribution = payout + rake", "多池/平分/退款守恒"),
        ("Lifecycle", "13 / 13", "每态恢复 + SETTLED->WAITING"),
    ], [2500, 2600, 4260])

    add_heading(document, "10. 机器人与玩家输赢控制", 1)
    add_paragraph(document, "Phase A 只有一个认证真人。其余五席由服务端机器人持有私牌和决策；客户端只看到允许公开的牌背、动作和摊牌信息。")
    add_table(document, ["能力", "实现", "安全边界"], [
        ("机器人节奏", "SOURCE_DELAYED 1-5 秒", "deadline/节奏进入快照"),
        ("机器人生命周期", "5-200 秒", "局间换人；Out/PlayerAct"),
        ("动作合法性", "与真人共用 legal action validator", "机器人不能绕过规则"),
        ("完整 deal commitment", "开牌前一次提交", "客户端拿不到 deck/candidate"),
        ("OFF", "一次正常洗牌与固定公牌", "不按目标选底牌"),
        ("POSITIVE/NEGATIVE", "在同一剩余牌堆分配候选底牌", "作用于目标真人"),
        ("target stop", "按 terminal net 幂等推进", "达到目标后下一手 OFF"),
    ], [2100, 3900, 3360], font_size=8.6)
    control = acceptance["controlConfidence95"]
    add_table(document, ["方向", "平均 net", "95% 下界", "95% 上界"], [
        (direction, values["averageNet"], round(values["lower95"], 2), round(values["upper95"], 2))
        for direction, values in control.items()
    ], [1600, 2500, 2500, 2760])
    add_callout(document, "面试重点", "控牌不是前端“改结果”。服务端先固定完整牌局，再投影 viewer-safe 事件；控制进度按真实 terminal net 更新，并有 OFF、非目标隔离和目标停止。", fill="EEF4FA")

    add_heading(document, "11. 前端效果与实现", 1)
    add_figure(document, popup_flow, "图 5  Rule、Set、Night 与 844×390")
    add_figure(document, settlement_flow, "图 6  原版自然结算；边池/退款规则保留为服务端逻辑证据")
    document.add_page_break()
    add_table(document, ["效果", "原资源/API", "触发"], [
        ("Load", "Load Spine start -> idle + sound/back", "authenticated session 后"),
        ("Room 入场", "top/content tween + 原人物/房间卡", "GameSessions"),
        ("发底牌", "NodePool + move/rotate/scale", "FaCards"),
        ("下注筹码", "chip NodePool + moveTo", "ActBet"),
        ("收池", "4 chips -> label/allbet", "PublicCards/Result"),
        ("公牌", "牌背飞入 -> 翻正面", "flop/turn/river"),
        ("牌型", "typeImg SpriteAtlas", "px / Result value"),
        ("摊牌", "win/lose/bigwin 原模板", "Result cards"),
        ("派奖", "winlabel + per-seat Spine", "winner Label active"),
        ("边池/退款", "动态克隆原 allbet row", "pots / uncalledReturns"),
        ("昼夜", "原 Night 黑色 Splash opacity 40", "Set yj/bt/zd"),
    ], [1800, 3800, 3760], font_size=8.5)

    document.add_page_break()
    add_heading(document, "12. 刷新、重连与安全", 1)
    add_numbered(document, [
        "launchCode 只交换一次；成功后立即从 current URL 删除。",
        "标签页 sessionStorage 只保留 opaque sessionToken、sessionId、roomId、roomLevel。",
        "刷新时 sessionToken 重新 context/init；不从客户端恢复牌局。",
        "Cocos 以 roomId 走原 reconnect Load/Room/Main，FinishLoad 请求 Redis snapshot。",
        "WS 断线采用 5s 起步指数退避；成功后归零，避免重启风暴。",
        "后端进程重启后，同 session/room/viewer/6席保持；Redis CAS/AOF 测试通过。",
        "退出游戏清除 tab-scoped reconnect key。",
    ])
    add_table(document, ["客户端允许", "客户端禁止"], [
        ("viewer uid/seat/stack", "complete deck / candidate deals"),
        ("本人私牌、公开公牌", "机器人未公开私牌"),
        ("合法动作按钮", "赢家、payout、control direction"),
        ("短期 opaque sessionToken", "appSecret、WalletGateway、RoundService"),
        ("roomId/roomLevel", "authority revision 作为提交权威"),
    ], [4680, 4680])

    add_heading(document, "13. 问题、排查与决策记录", 1)
    decision_rows = [
        ("旧 Phase A 是自绘 UI", "源码/截图/类名审计", "彻底作废，独立工程重新物化原 Prefab", "学习目标要求原版"),
        ("Creator build 过但 runtime missing", "浏览器报 __spreadArrays / missing class", "恢复 Hall TypeScript helpers 为 Creator plugin", "必须早于 bundle scripts"),
        ("语义组件找不到 DZPK module", "Main bundle 早于 DZPK bundle", "bundle load 后再挂接；最终直接语义组件", "不靠外部网页壳"),
        ("UID 不一致", "RoomInfo 找不到 viewer", "Hall_Connect 顶层 uid 回写 GameContext", "服务端身份优先"),
        ("试玩余额不足", "原 min_gold 20万-100万", "dev trial top-up 到 100万", "不改原房间门槛"),
        ("现有本地 DB ledger 不安全", "0000 hash mismatch", "新建 task DB，43 migrations 前滚", "不篡改旧 ledger"),
        ("刷新丢会话", "launchCode 已清除", "tab sessionToken + room identity；server snapshot", "不存 event/bot/deck"),
        ("重启刷 WS error", "每秒直连", "5s 指数退避 + 原子重启测试", "减少故障风暴"),
        ("Set 缺 show_day_night", "真实点击报错", "恢复原 Night 节点和 Cocos 内置方法", "按用户要求不套 Hall"),
        ("结果时按钮重现", "close raise 恢复 bet", "区分关闭加注与全局隐藏", "方法职责明确"),
        ("退款 BitmapFont 缺中文字", "win.png 无退/回 glyph", "金额保持数字，语义放原 pot row", "零缺字、保留原字体"),
        ("200 手测试默认超时", "fresh review 下 case 超过 Bun 5s", "局部预算 30s；同命令重跑 67/67", "保留 200 手覆盖，不放宽全局"),
    ]
    add_table(document, ["问题", "证据", "决策", "理由"], decision_rows, [2100, 2300, 2900, 2060], font_size=7.9)

    add_heading(document, "14. 构建、调试与验收", 1)
    add_heading(document, "14.1 常用命令", 2)
    add_code_block(document, r"""
# 重新从权威源物化并重绑 semantic components
scripts\materialize-original-source.ps1

# 边界与精确 Creator 2.4.7 build
scripts\verify-original-client-boundary.ps1
scripts\build-creator247.ps1

# 预览
scripts\start-preview.ps1 -Port 3002

# GameHub TRIAL launch（需要本地商户显式允许 dzpk-955）
pnpm kg-micro-shell:play-url -- dzpk --trial --no-stage --merchant=...

# Engine / Backend
pnpm --filter @gamehub/kg-dzpk-engine test
pnpm --filter @gamehub/kg-dzpk-engine test:acceptance
bun test apps/backend/src/modules/games/dzpk ...
""")
    add_heading(document, "14.2 最终验证数字", 2)
    add_table(document, ["门", "结果", "证据"], [
        ("Creator fresh import/build", "Passed", "build-creator247.ps1"),
        ("Boundary", "6 Prefab / 5 semantic core / no authority leak", "verify-original-client-boundary.ps1"),
        ("Engine focused", "33 pass / 0 fail / 141 expects", "5 test files"),
        ("Backend focused", "67 review / 130 expanded; 0 fail", "530 / 1,206 expects；默认 timeout"),
        ("5-card", f"{acceptance['fiveCard']['combinations']:,}", "exhaustive"),
        ("7-card", f"{acceptance['sevenCard']['samples']:,} / 0 mismatch", "independent oracle"),
        ("Bot soak", f"{acceptance['botSoak']['hands']:,} hands / {acceptance['botSoak']['playerActions']:,} actions", "50,000 unique deals"),
        ("Natural browser", "2 distinct hands / board 3,4,5", "original UI buttons"),
        ("Refresh / restart", "same session/room/viewer/6 seats", "0 errors / 0 warnings"),
        ("Rule/Set/Night", "actual clicks passed", "night 255 / day 0"),
        ("844x390", "Canvas exactly 844x390", "mobile screenshot"),
    ], [2200, 3000, 4160], font_size=8.5)

    add_heading(document, "15. 面试问答与简历表述", 1)
    interview_questions = [
        ("为什么不直接复用原 JS？", "原文件保留作证据，但变量和职责不可维护。通过相同 Prefab 位置挂语义组件，兼得画面 parity 与可读源码。"),
        ("Prefab 和 .meta 为什么重要？", "Prefab 保存节点/组件序列化；.meta UUID 连接 SpriteFrame、Atlas、Spine、脚本 Class ID。丢 UUID 会出现 missing asset/script。"),
        ("Asset Bundle 怎么工作？", "Boot 先 loadBundle('DZPK')，之后 bundle.load prefab；bundle script 注册后才能实例化其组件。"),
        ("为什么 build 成功仍会 runtime 失败？", "构建只证明可编译。缺全局 helper、模块加载时序、序列化属性、Hall 服务函数都要靠浏览器发现。"),
        ("Controller 和 Presentation 怎么分工？", "Controller 消费 source event、更新 model、排动画队列；Presentation 只操作原节点和动画。"),
        ("重连怎么保证一致？", "sessionToken 重新鉴权，roomId 触发原 reconnect path，Redis versioned snapshot 是事实源；不回放客户端旧动画日志。"),
        ("如何避免重复 Result？", "Result 携带 handId/revision/actionSeq；Controller fingerprint 幂等，后端 request_id/CAS 防重复写。"),
        ("边池怎么算？", "按不同 contribution tier 构造层，每层 eligibility 独立；派奖+退款+rake 与总贡献守恒。"),
        ("short all-in 为什么不重开加注？", "未达到上一次 full raise size 时只增加 toCall，不给已行动玩家新的 raise right。"),
        ("为什么花色不破平？", "标准德州只比较五张牌 rank/kicker；花色无大小。"),
        ("控牌放哪里？", "服务端 deal commitment 边界；客户端没有 deck/candidate/control facts。"),
        ("如何证明控牌方向？", "OFF/POSITIVE/NEGATIVE 各 1000 手，95% 区间分离；进度按真实 terminal net 幂等推进。"),
        ("机器人为什么也走合法动作？", "避免 bot 形成旁路，所有动作统一验证 hand/seat/revision/legal action。"),
        ("为什么 Phase A 不做真钱？", "德州是一手多街，不能把每次 call/raise 当即时 round；需要后续 staged hand economics。"),
        ("为什么 sessionStorage 可以接受？", "只存 tab-scoped opaque session token 和导航身份，用于刷新；不存牌局/事件/deck，退出清除。生产可继续升级 cookie/origin 方案。"),
        ("如何测低概率边池/平分？", "逻辑用 deterministic unit/acceptance；视觉用断开 WS 后的 forced viewer-safe Result，并明确不当自然概率证据。"),
        ("844×390 怎么适配？", "1334×750 设计分辨率 + SHOW_ALL + 原 Widget；真实 viewport/canvas 均为 844×390。"),
        ("Set 的昼夜怎么补？", "从原 Main.fire 恢复 Night 黑色 Splash 节点，Cocos 内实现 show_day_night，opacity 255/0。"),
        ("最大的排障收获？", "源码身份、构建、运行、协议、视觉、恢复是不同证据；每一层都需要独立 gate。"),
        ("下一步 3.8 难点？", "2.x API/序列化/Bundle/Action 到 3.8 的模块、装饰器、资源和构建变化；必须先冻结 2.4.7 golden parity。"),
    ]
    for question_index, (question, answer) in enumerate(interview_questions, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        question_run = paragraph.add_run(f"Q{question_index}. {question}")
        set_run_font(question_run, bold=True, color=DARK_BLUE, size=10.5)
        answer_paragraph = document.add_paragraph()
        answer_paragraph.paragraph_format.left_indent = Inches(0.18)
        answer_paragraph.paragraph_format.space_after = Pt(6)
        answer_run = answer_paragraph.add_run("A. " + answer)
        set_run_font(answer_run, size=10.2, color=INK)

    add_heading(document, "15.1 简历可据实表述", 2)
    add_callout(document, "项目名称", "KG 德州扑克原版 Cocos 客户端还原与 GameHub TRIAL 接入（Creator 2.4.7）", fill="EEF4FA")
    add_bullets(document, [
        "基于 KG Cocos/PHP 权威源码完成 6 个 Prefab、695 个资源文件和原事件协议审计，恢复 Load->Room->DZPKMain 原运行链。",
        "将 5 个含糊核心脚本重构为可维护的 Cocos 语义组件，保留原节点、SpriteAtlas、Spine、Particle、Animation 与交互时序。",
        "实现服务端权威德州 engine：best-five、主/边池、平分/odd chip、未跟注退回、full raise/short all-in、幂等与 13 状态快照。",
        "实现 1 真人 + 5 机器人、1-5 秒动作节奏、5-200 秒生命周期、Redis CAS/AOF 重启恢复和 15 秒自动下一手。",
        "接入 GameHub context/init + authenticated KG WS，完成 viewer privacy、source uid 绑定、刷新/进程重启与 TRIAL/REAL 双重门。",
        "完成 259 万五张穷举、100 万七张零差异、5 万手 115 万动作 soak，并用真实 Chromium 验收两手、Rule/Set/Night 与 844×390。",
    ])
    add_callout(document, "不要写", "“完成 Creator 3.8、2-6 真人、真钱闭环、正式生产上线”目前都不属实。面试时以代码、测试和取舍讲清楚，比夸大范围更可信。", fill=RISK_FILL, accent=RGBColor(155, 28, 28))

    add_heading(document, "16. 难度评估与后续路线", 1)
    add_table(document, ["目标", "难度", "理由"], [
        ("只恢复素材/能打开 Canvas", "3/10", "不能称原版源码还原"),
        ("原 Prefab + 语义源码 + 离线牌局", "8/10", "旧 Hall 依赖、模块时序、规则和动画"),
        ("本阶段 GameHub TRIAL", "8.5/10", "认证、Redis、viewer、重连、测试矩阵"),
        ("Creator 3.8 迁移", "8.5-9/10", "序列化/API/Bundle/构建变化"),
        ("2-6 真人 + 真钱", "9.5/10", "并发、断线、hand economics、审计/对账"),
    ], [3300, 1400, 4660])
    add_heading(document, "16.1 下一阶段建议顺序", 2)
    add_numbered(document, [
        "冻结 2.4.7 golden screenshots、事件 trace、Prefab hashes 和 full acceptance。",
        "建立 Creator 3.8 空工程与 API compatibility ledger；先不改业务。",
        "逐 Prefab 迁移资源/序列化组件，保持 source event 和服务端不变。",
        "恢复 3.8 Web build、桌面/844×390 视觉和两手主循环。",
        "再扩 2-6 认证真人：seat admission、per-viewer projection、disconnect/rejoin。",
        "最后设计 REAL staged hand economics、wallet/order/ledger/audit/admin reconciliation。",
    ])

    add_heading(document, "附录 A. 关键身份与证据", 1)
    add_table(document, ["项目", "值"], [
        ("Source commit", "b5694d576c482e02dc00a33f51eea633b9cd647f"),
        ("Source game", "KG 119 / client 19 / GameHub dzpk-955"),
        ("Creator", "2.4.7"),
        ("DZPK files", "695 / 10,781,546 authority bytes"),
        ("Main nodes", "321"),
        ("Semantic controller", "UUID 6b8f401e-0cd5-4e75-a424-48316c0f02e7"),
        ("Semantic presentation", "UUID 2de88e3d-5d6c-47b1-b942-913966e6ac3f"),
        ("Acceptance manifest", acceptance["manifestId"]),
        ("Source math", acceptance["sourceMathVersion"]),
        ("Bank", "SourcePresentButGameRuntimeNotApplicable"),
    ], [2600, 6760])

    add_heading(document, "附录 B. 证据文档索引", 1)
    add_bullets(document, [
        "docs/phase-a-original-cocos-parity-audit.md",
        "docs/phase-a-cocos-naming-mapping.md",
        "docs/phase-a-cocos-component-mapping.md",
        "docs/phase-a-source-function-protocol-matrix.md",
        "docs/phase-a-validation-evidence.md + phase-a-fresh-independent-review.md",
        "docs/evidence/phase-a-acceptance-final.json",
    ])
    final_reminder = add_paragraph(
        document,
        "最终提醒：面试前请亲自打开 Prefab、断点跟事件并至少手动玩两手；重点是能沿源码解释失败、取舍与证据。",
        bold_prefix="最终提醒：",
        after=0,
    )
    final_reminder.paragraph_format.space_before = Pt(6)
    final_shading = OxmlElement("w:shd")
    final_shading.set(qn("w:fill"), "EEF4FA")
    final_reminder._p.get_or_add_pPr().append(final_shading)

    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    built_path = build_document()
    print(json.dumps({"docx": str(built_path)}, ensure_ascii=False))
